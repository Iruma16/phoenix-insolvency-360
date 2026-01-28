"""
Parser de archivos Word (.docx) para Phoenix Legal.

FASE 1D: MULTI-FORMATO
Objetivo: Extraer contenido de archivos Word de forma estructurada.

Casos de uso:
- Informes previos de auditores
- Contratos mercantiles
- Escritos de acreedores
- Comunicaciones formales
- Documentos societarios

PRINCIPIOS:
- Extraer todo el texto visible (párrafos + tablas)
- Preservar estructura básica
- Detectar secciones/párrafos
- No interpretar ni analizar
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument


class WordParseResult:
    """
    Resultado del parsing de un archivo Word.

    Atributos:
        texto: Texto extraído (representación textual del contenido)
        num_paginas: Estimado de páginas (basado en párrafos, aproximado)
        tipo_documento: Siempre "word"
        page_offsets: Diccionario con offsets estimados por sección
        paragraphs_content: Contenido por párrafo (opcional, para análisis)
        tables_content: Contenido de tablas (opcional)
    """

    def __init__(
        self,
        texto: str,
        num_paginas: int,
        page_offsets: dict[int, tuple[int, int]],
        paragraphs_content: Optional[list[str]] = None,
        tables_content: Optional[list[list[list[str]]]] = None,
    ):
        self.texto = texto
        self.num_paginas = num_paginas
        self.tipo_documento = "word"
        self.page_offsets = page_offsets
        self.paragraphs_content = paragraphs_content or []
        self.tables_content = tables_content or []


def estimate_pages_from_paragraphs(num_paragraphs: int) -> int:
    """
    Estima el número de páginas basándose en el número de párrafos.

    Heurística conservadora:
    - ~20-25 párrafos por página
    - Mínimo 1 página

    Args:
        num_paragraphs: Número de párrafos en el documento

    Returns:
        Estimado de páginas (mínimo 1)
    """
    if num_paragraphs == 0:
        return 1

    # Asumimos ~22 párrafos por página (promedio)
    estimated = max(1, (num_paragraphs + 21) // 22)
    return estimated


def parse_word_file(file_path: str) -> WordParseResult:
    """
    Parsea un archivo Word y extrae su contenido como texto.

    Args:
        file_path: Ruta al archivo Word (.docx)

    Returns:
        WordParseResult con el contenido extraído

    Raises:
        Exception: Si el archivo no se puede leer o no es un Word válido

    Estrategia de extracción:
        1. Abrir documento con python-docx
        2. Extraer párrafos (con estilos si están presentes)
        3. Extraer tablas
        4. Intercalar párrafos y tablas en orden
        5. Generar representación textual
        6. Estimar páginas y calcular offsets
    """
    # Abrir documento
    doc = DocxDocument(file_path)

    full_text_parts: list[str] = []
    paragraphs_content: list[str] = []
    tables_content: list[list[list[str]]] = []
    page_offsets: dict[int, tuple[int, int]] = {}

    current_offset = 0
    paragraph_count = 0

    # Título del documento
    doc_header = f"\n{'=' * 80}\n"
    doc_header += "DOCUMENTO WORD\n"
    doc_header += f"{'=' * 80}\n\n"
    full_text_parts.append(doc_header)
    current_offset += len(doc_header)

    # Iterar sobre todos los elementos del documento (párrafos y tablas)
    # Nota: python-docx no preserva el orden exacto de tablas/párrafos
    # Primero extraemos párrafos, luego tablas

    # 1. PÁRRAFOS
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:  # Solo incluir párrafos no vacíos
            paragraph_count += 1
            paragraphs_content.append(text)

            # Detectar si es título/encabezado (por estilo)
            style = paragraph.style.name if paragraph.style else ""
            prefix = ""

            if "Heading" in style or "Título" in style:
                prefix = f"[{style}] "
                text_line = f"\n{prefix}{text}\n"
            else:
                text_line = f"{text}\n"

            full_text_parts.append(text_line)
            current_offset += len(text_line)

    # 2. TABLAS
    if doc.tables:
        separator = f"\n{'=' * 80}\n"
        separator += "TABLAS\n"
        separator += f"{'=' * 80}\n\n"
        full_text_parts.append(separator)
        current_offset += len(separator)

        for table_idx, table in enumerate(doc.tables):
            table_header = f"\n--- TABLA {table_idx + 1} ---\n"
            full_text_parts.append(table_header)
            current_offset += len(table_header)

            table_data: list[list[str]] = []

            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                table_data.append(row_values)

                # Añadir fila al texto (formato tabular)
                row_text = " | ".join(row_values)
                row_line = f"{row_text}\n"
                full_text_parts.append(row_line)
                current_offset += len(row_line)

            tables_content.append(table_data)

            # Separador entre tablas
            table_sep = "\n"
            full_text_parts.append(table_sep)
            current_offset += len(table_sep)

    # Estimar páginas basándose en número de párrafos
    estimated_pages = estimate_pages_from_paragraphs(paragraph_count)

    # Calcular offsets estimados por "página"
    # Dividimos el texto total en N secciones aproximadas
    total_length = current_offset
    chars_per_page = total_length // estimated_pages if estimated_pages > 0 else total_length

    for page_num in range(estimated_pages):
        start = page_num * chars_per_page
        end = min((page_num + 1) * chars_per_page, total_length)
        page_offsets[page_num] = (start, end)

    # Unir todo el texto
    full_text = "".join(full_text_parts)

    return WordParseResult(
        texto=full_text,
        num_paginas=estimated_pages,
        page_offsets=page_offsets,
        paragraphs_content=paragraphs_content,
        tables_content=tables_content,
    )


def parse_word_stream(file_stream: io.BytesIO, filename: str) -> WordParseResult:
    """
    Parsea un archivo Word desde un stream de bytes.

    Útil para archivos subidos vía API sin guardar temporalmente en disco.

    Args:
        file_stream: Stream de bytes con el contenido del Word
        filename: Nombre original del archivo (solo para logging)

    Returns:
        WordParseResult con el contenido extraído
    """
    # python-docx puede leer directamente desde BytesIO
    doc = DocxDocument(file_stream)

    full_text_parts: list[str] = []
    paragraphs_content: list[str] = []
    tables_content: list[list[list[str]]] = []
    page_offsets: dict[int, tuple[int, int]] = {}

    current_offset = 0
    paragraph_count = 0

    # Título del documento
    doc_header = f"\n{'=' * 80}\n"
    doc_header += f"DOCUMENTO WORD: {filename}\n"
    doc_header += f"{'=' * 80}\n\n"
    full_text_parts.append(doc_header)
    current_offset += len(doc_header)

    # 1. PÁRRAFOS
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:  # Solo incluir párrafos no vacíos
            paragraph_count += 1
            paragraphs_content.append(text)

            # Detectar si es título/encabezado (por estilo)
            style = paragraph.style.name if paragraph.style else ""
            prefix = ""

            if "Heading" in style or "Título" in style:
                prefix = f"[{style}] "
                text_line = f"\n{prefix}{text}\n"
            else:
                text_line = f"{text}\n"

            full_text_parts.append(text_line)
            current_offset += len(text_line)

    # 2. TABLAS
    if doc.tables:
        separator = f"\n{'=' * 80}\n"
        separator += "TABLAS\n"
        separator += f"{'=' * 80}\n\n"
        full_text_parts.append(separator)
        current_offset += len(separator)

        for table_idx, table in enumerate(doc.tables):
            table_header = f"\n--- TABLA {table_idx + 1} ---\n"
            full_text_parts.append(table_header)
            current_offset += len(table_header)

            table_data: list[list[str]] = []

            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                table_data.append(row_values)

                # Añadir fila al texto (formato tabular)
                row_text = " | ".join(row_values)
                row_line = f"{row_text}\n"
                full_text_parts.append(row_line)
                current_offset += len(row_line)

            tables_content.append(table_data)

            # Separador entre tablas
            table_sep = "\n"
            full_text_parts.append(table_sep)
            current_offset += len(table_sep)

    # Estimar páginas basándose en número de párrafos
    estimated_pages = estimate_pages_from_paragraphs(paragraph_count)

    # Calcular offsets estimados por "página"
    total_length = current_offset
    chars_per_page = total_length // estimated_pages if estimated_pages > 0 else total_length

    for page_num in range(estimated_pages):
        start = page_num * chars_per_page
        end = min((page_num + 1) * chars_per_page, total_length)
        page_offsets[page_num] = (start, end)

    # Unir todo el texto
    full_text = "".join(full_text_parts)

    return WordParseResult(
        texto=full_text,
        num_paginas=estimated_pages,
        page_offsets=page_offsets,
        paragraphs_content=paragraphs_content,
        tables_content=tables_content,
    )


def detect_word_type(filename: str) -> Optional[str]:
    """
    Detecta si un archivo es un Word válido por su extensión.

    Args:
        filename: Nombre del archivo

    Returns:
        "docx" si es Word, None si no lo es
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        return "docx"
    elif ext == ".doc":
        # .doc es formato antiguo (binario), python-docx solo soporta .docx
        # Podríamos intentar leerlo pero probablemente falle
        return "doc_legacy"
    else:
        return None
