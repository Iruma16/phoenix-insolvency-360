from datetime import datetime, timezone
from io import BytesIO


def generate_docx_report(legal_report, case) -> bytes:
    """
    Genera informe en formato Word (.docx) editable CON TRAZABILIDAD LEGAL.

    CRÍTICO: Word debe incluir:
    - Metadata de trazabilidad (IDs, hashes, versiones)
    - Referencias exactas a chunks (página, offsets)
    - Hash de integridad del contenido
    - Información suficiente para auditoría legal

    Args:
        legal_report: LegalReport Pydantic model
        case: Case SQLAlchemy model

    Returns:
        bytes: Contenido del archivo DOCX
    """
    try:
        import hashlib
        import json
        import uuid

        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt, RGBColor

        doc = Document()

        # Estilos de documento
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # ==========================================
        # PORTADA
        # ==========================================
        title = doc.add_heading("INFORME LEGAL PRELIMINAR", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f"Caso: {case.name}")
        doc.add_paragraph(f"ID Caso: {legal_report.case_id}")
        doc.add_paragraph(f"Sistema: {legal_report.source_system} v{legal_report.schema_version}")
        doc.add_paragraph(f"Generado: {legal_report.generated_at.strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()

        # Disclaimer
        disclaimer_para = doc.add_paragraph()
        disclaimer_run = disclaimer_para.add_run(
            "AVISO LEGAL: Este informe es un análisis técnico preliminar generado automáticamente. "
            "NO constituye asesoramiento legal ni dictamen jurídico. Requiere revisión y validación por "
            "profesional legal cualificado antes de tomar decisiones."
        )
        disclaimer_run.font.color.rgb = RGBColor(220, 38, 38)
        disclaimer_run.bold = True

        doc.add_page_break()

        # ==========================================
        # METADATA DE TRAZABILIDAD (CRÍTICO)
        # ==========================================
        doc.add_heading("METADATA DE TRAZABILIDAD", 1)

        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(
            "Esta sección es OBLIGATORIA para auditoría legal. "
            "Permite verificar la integridad y trazabilidad del informe."
        )
        meta_run.font.size = Pt(9)
        meta_run.italic = True

        # ✅ Generar Report ID único (UUID)
        report_id = str(uuid.uuid4())

        # ✅ Calcular hash FUERTE de integridad (incluye TODOS los campos críticos)
        content_for_hash = {
            "report_id": report_id,
            "case_id": legal_report.case_id,
            "issue_analyzed": legal_report.issue_analyzed,
            "generated_at": legal_report.generated_at.isoformat(),
            "findings": [],
        }

        # Incluir hallazgos con evidencias completas (orden determinista)
        findings_sorted = sorted(legal_report.findings, key=lambda f: f.statement)
        for f in findings_sorted:
            finding_data = {
                "statement": f.statement,
                "confidence_note": f.confidence_note if hasattr(f, "confidence_note") else None,
                "evidence": [],
            }

            # Incluir evidencias con offsets y ubicación (orden determinista)
            if hasattr(f, "evidence") and f.evidence:
                evidence_sorted = sorted(f.evidence[:10], key=lambda e: e.document_id)
                for ev in evidence_sorted:
                    ev_data = {
                        "document_id": ev.document_id,
                        "filename": ev.filename if hasattr(ev, "filename") else None,
                    }
                    if hasattr(ev, "location") and ev.location:
                        if hasattr(ev.location, "page_start"):
                            ev_data["page_start"] = ev.location.page_start
                        if hasattr(ev.location, "char_start"):
                            ev_data["char_start"] = ev.location.char_start
                        if hasattr(ev.location, "char_end"):
                            ev_data["char_end"] = ev.location.char_end
                    finding_data["evidence"].append(ev_data)

            content_for_hash["findings"].append(finding_data)

        # Hash determinista con JSON ordenado
        content_hash = hashlib.sha256(
            json.dumps(content_for_hash, sort_keys=True).encode()
        ).hexdigest()[:32]

        meta_table = doc.add_table(rows=6, cols=2)
        meta_table.style = "Light Grid Accent 1"

        meta_data = [
            ("Report ID", report_id),  # ✅ UUID único, no source_system
            ("Schema Version", legal_report.schema_version),
            ("Caso ID", legal_report.case_id),
            ("Generado", legal_report.generated_at.isoformat()),
            ("Hallazgos totales", str(len(legal_report.findings))),
            ("Content Hash", content_hash),
        ]

        for idx, (key, val) in enumerate(meta_data):
            row = meta_table.rows[idx]
            row.cells[0].text = key
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(val)

        doc.add_paragraph()

        # ==========================================
        # ASUNTO ANALIZADO
        # ==========================================
        doc.add_heading("1. ASUNTO ANALIZADO", 1)
        doc.add_paragraph(legal_report.issue_analyzed)

        # ==========================================
        # HALLAZGOS CON TRAZABILIDAD COMPLETA
        # ==========================================
        doc.add_heading(f"2. HALLAZGOS LEGALES ({len(legal_report.findings)})", 1)

        if not legal_report.findings:
            doc.add_paragraph(
                "No se identificaron hallazgos legales relevantes en el análisis técnico preliminar."
            )
        else:
            for idx, finding in enumerate(legal_report.findings, 1):
                # Título del hallazgo
                title_text = (
                    finding.statement[:80] + "..."
                    if len(finding.statement) > 80
                    else finding.statement
                )
                doc.add_heading(f"2.{idx}. {title_text}", 2)

                # Declaración completa
                doc.add_paragraph(finding.statement)

                # Nota de confianza
                if finding.confidence_note:
                    conf_para = doc.add_paragraph()
                    conf_run = conf_para.add_run(f"Nota de confianza: {finding.confidence_note}")
                    conf_run.font.size = Pt(9)
                    conf_run.font.color.rgb = RGBColor(107, 114, 128)

                # EVIDENCIAS CON TRAZABILIDAD COMPLETA
                if finding.evidence:
                    doc.add_heading(f"Evidencias con trazabilidad ({len(finding.evidence)}):", 3)

                    for ev_idx, evidence in enumerate(finding.evidence[:10], 1):  # Máx 10
                        # Texto principal
                        ev_filename = (
                            evidence.filename if evidence.filename else evidence.document_id
                        )
                        doc.add_paragraph(
                            f"{ev_idx}. Documento: {ev_filename}", style="List Bullet"
                        )

                        # ✅ TRAZABILIDAD DETALLADA con NULL-SAFETY (crítico para defensa legal)
                        trace_para = doc.add_paragraph(style="List Bullet 2")
                        trace_parts = []

                        # Document ID
                        trace_parts.append(f"Doc ID: {evidence.document_id}")

                        # Chunk ID si existe
                        if hasattr(evidence, "chunk_id") and evidence.chunk_id:
                            trace_parts.append(f"Chunk: {evidence.chunk_id}")

                        # ✅ NULL-SAFETY: Verificar que location existe
                        if hasattr(evidence, "location") and evidence.location:
                            # Ubicación de página
                            if (
                                hasattr(evidence.location, "page_start")
                                and evidence.location.page_start
                            ):
                                if (
                                    hasattr(evidence.location, "page_end")
                                    and evidence.location.page_end
                                    and evidence.location.page_end != evidence.location.page_start
                                ):
                                    trace_parts.append(
                                        f"Páginas: {evidence.location.page_start}-{evidence.location.page_end}"
                                    )
                                else:
                                    trace_parts.append(f"Página: {evidence.location.page_start}")
                            else:
                                trace_parts.append("Página: N/A")

                            # Offsets de carácter (crítico para verificación)
                            if (
                                hasattr(evidence.location, "char_start")
                                and evidence.location.char_start is not None
                            ):
                                char_end = (
                                    evidence.location.char_end
                                    if hasattr(evidence.location, "char_end")
                                    and evidence.location.char_end is not None
                                    else "N/A"
                                )
                                trace_parts.append(
                                    f"Offsets: {evidence.location.char_start}-{char_end}"
                                )
                            else:
                                trace_parts.append("Offsets: N/A")

                            # Método de extracción
                            if (
                                hasattr(evidence.location, "extraction_method")
                                and evidence.location.extraction_method
                            ):
                                trace_parts.append(f"Método: {evidence.location.extraction_method}")
                            else:
                                trace_parts.append("Método: N/A")
                        else:
                            # ✅ Fallback estructurado parseable (mantiene formato consistente)
                            trace_parts.append("Página: N/A | Offsets: N/A | Método: N/A")

                        trace_run = trace_para.add_run(" | ".join(trace_parts))
                        trace_run.font.size = Pt(8)
                        trace_run.font.color.rgb = RGBColor(75, 85, 99)

                    if len(finding.evidence) > 10:
                        more_para = doc.add_paragraph(
                            f"... y {len(finding.evidence) - 10} evidencia(s) más",
                            style="List Bullet",
                        )
                        more_para.runs[0].italic = True

        # ==========================================
        # PIE DE PÁGINA CON HASH
        # ==========================================
        doc.add_page_break()
        doc.add_heading("FIN DEL INFORME", 1)

        footer_para = doc.add_paragraph(
            f"Sistema: {legal_report.source_system} | Generado: {datetime.now(timezone.utc).strftime('%d/%m/%Y %H:%M:%S UTC')}"
        )
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(107, 114, 128)

        # Hash de integridad
        hash_para = doc.add_paragraph(f"Content Hash (verificación): {content_hash}")
        hash_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hash_run = hash_para.runs[0]
        hash_run.font.size = Pt(8)
        hash_run.font.color.rgb = RGBColor(156, 163, 175)

        # Guardar en BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    except ImportError:
        raise RuntimeError("python-docx no está instalado. Instala con: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Error generando informe Word: {str(e)}")
