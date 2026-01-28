from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any, Optional, Union

import pdfplumber
from docx import Document as DocxDocument
from openai import OpenAI

# FASE 2A: Parser de Email
from app.services.email_parser import parse_eml_stream, parse_msg_stream

# FASE 1C: Parser de Excel
from app.services.excel_parser import ExcelParseResult, parse_excel_stream

# FASE 2A: Parser OCR
from app.services.ocr_parser import (
    is_tesseract_available,
    ocr_pdf_from_bytes,
    should_apply_ocr_to_pdf,
)

# FASE 1D: Parser de Word
from app.services.word_parser import WordParseResult, parse_word_stream

# =========================================================
# INICIALIZACI√ìN
# =========================================================

print("üü¢ [INGESTA] Inicializando m√≥dulo de ingesta")

# Lazy initialization: cliente OpenAI se crea solo cuando se necesita
_openai_client: Optional[OpenAI] = None

if TYPE_CHECKING:
    import pandas as pd


def _get_openai_client() -> OpenAI:
    """
    Obtiene cliente OpenAI (lazy initialization).
    Solo se inicializa cuando se necesita realmente.
    """
    global _openai_client
    if _openai_client is None:
        from app.core.config import settings

        # Solo crear cliente si hay API key configurada
        if settings.openai_api_key:
            _openai_client = OpenAI(api_key=settings.openai_api_key)
        else:
            # Si no hay API key, crear cliente sin key (fallar√° al usarse, pero no al importar)
            _openai_client = OpenAI()
    return _openai_client


# =========================================================
# DATACLASS PARA RESULTADO DE PARSING
# =========================================================


@dataclass
class ParsingResult:
    """
    Resultado del parsing de un documento.
    Incluye texto extra√≠do y metadatos para validaci√≥n.
    """

    texto: str
    num_paginas: int  # N√∫mero de p√°ginas detectadas
    tipo_documento: str  # pdf, docx, txt, etc.
    page_offsets: dict[int, tuple[int, int]] | None = None  # {page_num: (start_char, end_char)}
    structured_data: dict[str, Any] | None = None  # FASE 2A: Datos estructurados (facturas, etc.)
    # FASE 2A: Metadata de OCR (trazabilidad legal)
    ocr_applied: bool = False
    ocr_pages: list[int] | None = None  # P√°ginas procesadas con OCR
    ocr_language: str | None = None  # Idioma usado (spa, eng)
    ocr_chars_detected: int | None = None  # Caracteres detectados por OCR


# =========================================================
# UTILIDADES GENERALES
# =========================================================


def normalizar_fecha(fecha_str: str) -> Optional[str]:
    """
    Intenta convertir m√∫ltiples formatos de fecha a YYYY-MM-DD.
    """
    if not fecha_str:
        return None

    formatos = [
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%d-%m-%Y",
        "%d.%m.%Y",
        "%Y/%m/%d",
    ]

    for fmt in formatos:
        try:
            fecha = datetime.strptime(fecha_str.strip(), fmt).strftime("%Y-%m-%d")
            print(f"‚úÖ [FECHA] Normalizada '{fecha_str}' ‚Üí '{fecha}'")
            return fecha
        except ValueError:
            continue

    print(f"‚ö†Ô∏è [FECHA] No se pudo normalizar: {fecha_str}")
    return fecha_str


# =========================================================
# INGESTA PDF / TXT
# =========================================================


def leer_pdf(file_stream) -> ParsingResult:
    """
    Lee un archivo PDF y extrae todo el texto.
    Soporta tanto rutas de archivo (string) como streams.
    Retorna ParsingResult con texto y metadatos.

    CORRECCI√ìN: Calcula page_offsets para trazabilidad de p√°ginas.
    FASE 2A: Si PDF es escaneado, aplica OCR autom√°ticamente.
    """
    print("üìÑ [PDF] Inicio lectura de PDF")
    texto_completo = ""
    num_paginas = 0
    page_offsets = {}  # {page_num: (start_char, end_char)}

    try:
        # Leer bytes del stream para poder reutilizarlo
        if hasattr(file_stream, "read"):
            pdf_bytes = file_stream.read()
        else:
            with open(file_stream, "rb") as f:
                pdf_bytes = f.read()

        # Intentar extracci√≥n normal primero
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            num_paginas = len(pdf.pages)
            print(f"üìÑ [PDF] N√∫mero de p√°ginas: {num_paginas}")

            current_offset = 0
            for i, page in enumerate(pdf.pages):
                page_num = i + 1  # P√°ginas comienzan en 1
                page_start = current_offset

                text = page.extract_text()
                if text:
                    texto_completo += text + "\n"
                    current_offset += len(text) + 1  # +1 por el \n
                else:
                    print(f"‚ö†Ô∏è [PDF] P√°gina {i+1} sin texto")

                page_end = current_offset
                page_offsets[page_num] = (page_start, page_end)

        # FASE 2A: Detectar si es PDF escaneado y aplicar OCR
        ocr_applied = False
        ocr_pages = None
        ocr_language = None
        ocr_chars_detected = None

        if should_apply_ocr_to_pdf(pdf_bytes, threshold_chars=50):
            print("üîç [PDF] PDF escaneado detectado, aplicando OCR...")

            if is_tesseract_available():
                ocr_result = ocr_pdf_from_bytes(pdf_bytes, lang="spa")

                if ocr_result.texto and len(ocr_result.texto) > len(texto_completo):
                    print(f"‚úÖ [PDF] OCR exitoso: {len(ocr_result.texto)} caracteres extra√≠dos")
                    texto_completo = ocr_result.texto
                    page_offsets = ocr_result.page_offsets
                    num_paginas = ocr_result.num_paginas
                    # Capturar metadata OCR
                    ocr_applied = True
                    ocr_pages = list(range(1, num_paginas + 1))
                    ocr_language = "spa"
                    ocr_chars_detected = len(texto_completo)
                else:
                    print("‚ö†Ô∏è [PDF] OCR no mejor√≥ extracci√≥n, usando texto original")
            else:
                print("‚ö†Ô∏è [PDF] Tesseract no disponible, no se puede aplicar OCR")
                print("‚ö†Ô∏è [PDF] Instalar: apt-get install tesseract-ocr tesseract-ocr-spa")

        print(f"‚úÖ [PDF] Texto extra√≠do ({len(texto_completo)} caracteres)")
        print(f"‚úÖ [PDF] Calculados offsets para {len(page_offsets)} p√°ginas")

        # FASE 2A: Detectar tipo de documento CON PRIORIDAD EXPL√çCITA
        # ORDEN: Accounting > Invoice > Legal > Generic
        structured_data = None
        from app.core.config import settings
        from app.services.accounting_parser import (
            is_financial_statement,
            parse_financial_statement_from_text,
        )
        from app.services.invoice_parser import is_likely_invoice, parse_invoice_from_text
        from app.services.legal_ner import extract_legal_entities, is_legal_document

        # PRIORIDAD 1: Estados financieros (m√°s espec√≠ficos)
        if is_financial_statement(texto_completo):
            print("üìä [PDF] Detectado estado financiero, extrayendo datos...")
            financial_stmt = parse_financial_statement_from_text(texto_completo)
            if financial_stmt:
                structured_data = {"financial_statement": financial_stmt.model_dump()}
                print(f"‚úÖ [PDF] Estado financiero: {financial_stmt.statement_type}")

        # PRIORIDAD 2: Facturas
        elif is_likely_invoice(texto_completo):
            print("üí∞ [PDF] Detectada posible factura, extrayendo datos...")
            invoice = parse_invoice_from_text(texto_completo)
            if invoice:
                structured_data = {"invoice": invoice.model_dump()}
                print(
                    f"‚úÖ [PDF] Factura extra√≠da: {invoice.invoice_number or 'N/A'} - {invoice.total_amount}‚Ç¨"
                )

        # PRIORIDAD 3: Documentos legales
        elif is_legal_document(texto_completo):
            if settings.enable_llm_extraction:  # ‚Üê Feature flag
                print("‚öñÔ∏è [PDF] Detectado documento legal, extrayendo entidades...")
                legal_doc = extract_legal_entities(texto_completo, use_llm=True)
                if legal_doc.entities:
                    structured_data = {"legal_document": legal_doc.model_dump()}
                    print(
                        f"‚úÖ [PDF] Documento legal: {legal_doc.document_type} con {len(legal_doc.entities)} entidades"
                    )
            else:
                print("‚ö†Ô∏è [PDF] LLM extraction deshabilitado, omitiendo extracci√≥n legal")

        return ParsingResult(
            texto=texto_completo,
            num_paginas=num_paginas,
            tipo_documento="pdf",
            page_offsets=page_offsets,
            structured_data=structured_data,
            ocr_applied=ocr_applied,
            ocr_pages=ocr_pages,
            ocr_language=ocr_language,
            ocr_chars_detected=ocr_chars_detected,
        )

    except Exception as e:
        print("‚ùå [PDF] Error leyendo PDF")
        print(f"‚ùå [PDF] Detalle: {e}")

        return ParsingResult(
            texto="",
            num_paginas=0,
            tipo_documento="pdf",
            page_offsets=None,
        )


def leer_txt(file_stream) -> ParsingResult:
    """
    Lee un archivo TXT.
    Soporta tanto rutas de archivo (string) como streams.
    Retorna ParsingResult con texto y metadatos.
    """
    print("üìÑ [TXT] Inicio lectura de TXT")
    try:
        # Si es una ruta (string o Path), abrir el archivo
        if isinstance(file_stream, (str, Path)):
            with open(file_stream, encoding="utf-8", errors="ignore") as f:
                texto = f.read()
        else:
            # Si es un stream, leerlo
            content = file_stream.read()
            texto = (
                content.decode("utf-8", errors="ignore")
                if isinstance(content, bytes)
                else str(content)
            )

        print(f"‚úÖ [TXT] Texto le√≠do ({len(texto)} caracteres)")

        # FASE 2A: Detectar tipo de documento CON PRIORIDAD EXPL√çCITA
        # ORDEN: Accounting > Invoice > Legal > Generic
        structured_data = None
        from app.core.config import settings
        from app.services.accounting_parser import (
            is_financial_statement,
            parse_financial_statement_from_text,
        )
        from app.services.invoice_parser import is_likely_invoice, parse_invoice_from_text
        from app.services.legal_ner import extract_legal_entities, is_legal_document

        # PRIORIDAD 1: Estados financieros (m√°s espec√≠ficos)
        if is_financial_statement(texto):
            print("üìä [TXT] Detectado estado financiero, extrayendo datos...")
            financial_stmt = parse_financial_statement_from_text(texto)
            if financial_stmt:
                structured_data = {"financial_statement": financial_stmt.model_dump()}
                print(f"‚úÖ [TXT] Estado financiero: {financial_stmt.statement_type}")

        # PRIORIDAD 2: Facturas
        elif is_likely_invoice(texto):
            print("üí∞ [TXT] Detectada posible factura, extrayendo datos...")
            invoice = parse_invoice_from_text(texto)
            if invoice:
                structured_data = {"invoice": invoice.model_dump()}
                print(
                    f"‚úÖ [TXT] Factura extra√≠da: {invoice.invoice_number or 'N/A'} - {invoice.total_amount}‚Ç¨"
                )

        # PRIORIDAD 3: Documentos legales
        elif is_legal_document(texto):
            if settings.enable_llm_extraction:  # ‚Üê Feature flag
                print("‚öñÔ∏è [TXT] Detectado documento legal, extrayendo entidades...")
                legal_doc = extract_legal_entities(texto, use_llm=True)
                if legal_doc.entities:
                    structured_data = {"legal_document": legal_doc.model_dump()}
                    print(
                        f"‚úÖ [TXT] Documento legal: {legal_doc.document_type} con {len(legal_doc.entities)} entidades"
                    )
            else:
                print("‚ö†Ô∏è [TXT] LLM extraction deshabilitado, omitiendo extracci√≥n legal")

        return ParsingResult(
            texto=texto,
            num_paginas=1,  # TXT no tiene concepto de p√°ginas
            tipo_documento="txt",
            structured_data=structured_data,
        )
    except Exception as e:
        print("‚ùå [TXT] Error leyendo TXT")
        print(f"‚ùå [TXT] Detalle: {e}")

        return ParsingResult(
            texto="",
            num_paginas=0,
            tipo_documento="txt",
        )


def leer_docx(file_stream, is_doc_legacy: bool = False) -> ParsingResult:
    """
    Lee un archivo DOCX y extrae todo el texto.
    Soporta tanto rutas de archivo (string) como streams.
    Retorna ParsingResult con texto y metadatos.

    Par√°metros
    ----------
    file_stream : str, Path, o stream
        Archivo DOCX a leer
    is_doc_legacy : bool
        Si True, indica que es un archivo .doc (formato antiguo).
        Se intentar√° leer pero puede fallar. Se mostrar√° un warning si falla.
    """
    file_type = "DOC (legacy)" if is_doc_legacy else "DOCX"
    tipo_doc = "doc" if is_doc_legacy else "docx"
    print(f"üìÑ [{file_type}] Inicio lectura de {file_type}")
    texto_completo = ""

    try:
        # Si file_stream es una ruta (string o Path), abrir el archivo
        if isinstance(file_stream, (str, Path)):
            doc = DocxDocument(str(file_stream))
        else:
            # Si es un stream, python-docx lo maneja directamente
            # Pero si es bytes, necesitamos crear un BytesIO
            from io import BytesIO

            if isinstance(file_stream, bytes):
                file_stream = BytesIO(file_stream)
            doc = DocxDocument(file_stream)

        # Extraer texto de todos los p√°rrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texto_completo += paragraph.text + "\n"

        # Extraer texto de las tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texto_completo += row_text + "\n"

        if is_doc_legacy and not texto_completo.strip():
            print(
                f"‚ö†Ô∏è  [{file_type}] Archivo .doc le√≠do pero sin contenido. Puede requerir conversi√≥n."
            )
        else:
            print(f"‚úÖ [{file_type}] Texto extra√≠do ({len(texto_completo)} caracteres)")

        return ParsingResult(
            texto=texto_completo.strip(),
            num_paginas=1,  # DOCX no tiene concepto de p√°ginas nativo
            tipo_documento=tipo_doc,
        )

    except Exception as e:
        if is_doc_legacy:
            print(f"‚ö†Ô∏è  [{file_type}] WARNING: No se pudo leer archivo .doc (formato antiguo)")
            print(
                f"‚ö†Ô∏è  [{file_type}] python-docx solo soporta .docx. Considera convertir a .docx con LibreOffice:"
            )
            print(
                f"‚ö†Ô∏è  [{file_type}]   libreoffice --headless --convert-to docx --outdir /tmp archivo.doc"
            )
            print(f"‚ö†Ô∏è  [{file_type}] Detalle del error: {e}")
        else:
            print("‚ùå [DOCX] Error leyendo DOCX")
            print(f"‚ùå [DOCX] Detalle: {e}")

        return ParsingResult(
            texto="",
            num_paginas=0,
            tipo_documento=tipo_doc,
        )


# =========================================================
# INGESTA CSV / EXCEL (BANCOS)
# =========================================================


def detectar_columna(columnas_disponibles, posibles_nombres):
    """
    Detecta autom√°ticamente columnas por sin√≥nimos.
    """
    cols_lower = [c.lower() for c in columnas_disponibles]

    for candidato in posibles_nombres:
        if candidato in cols_lower:
            return columnas_disponibles[cols_lower.index(candidato)]

    for col in columnas_disponibles:
        for candidato in posibles_nombres:
            if candidato in col.lower():
                return col

    return None


def normalizar_datos_banco(df: "pd.DataFrame") -> "pd.DataFrame":
    import pandas as pd

    print("üè¶ [BANCO] Normalizando datos bancarios")
    df.columns = df.columns.str.strip()
    cols = df.columns.tolist()
    print(f"üè¶ [BANCO] Columnas detectadas: {cols}")

    posibles_fechas = ["fecha", "date", "f.valor", "f.operacion", "d√≠a", "dia", "time"]
    posibles_conceptos = [
        "concepto",
        "descripcion",
        "detalle",
        "movimiento",
        "asunto",
        "transaccion",
        "transaction",
        "leyenda",
    ]
    posibles_importes = [
        "importe",
        "amount",
        "cantidad",
        "saldo",
        "euros",
        "valor",
        "cuantia",
        "monto",
    ]

    col_fecha = detectar_columna(cols, posibles_fechas)
    col_concepto = detectar_columna(cols, posibles_conceptos)
    col_importe = detectar_columna(cols, posibles_importes)

    print(
        f"üè¶ [BANCO] Mapeo columnas ‚Üí Fecha:{col_fecha}, Concepto:{col_concepto}, Importe:{col_importe}"
    )

    # Si no podemos mapear suficientes columnas t√≠picas de extracto bancario,
    # NO debemos ‚Äúrecortar‚Äù el DF a Fecha/Concepto/Importe porque destruir√≠a
    # CSVs que NO son extractos (p. ej. res√∫menes financieros con columnas propias).
    mapped = [c for c in (col_fecha, col_concepto, col_importe) if c]
    if len(mapped) < 2:
        print(
            "‚ö†Ô∏è [BANCO] No hay suficientes columnas bancarias mapeables; "
            "se devuelve el CSV/Excel sin normalizar."
        )
        return df

    nuevas_cols = {}
    if col_fecha:
        nuevas_cols[col_fecha] = "Fecha"
    if col_concepto:
        nuevas_cols[col_concepto] = "Concepto"
    if col_importe:
        nuevas_cols[col_importe] = "Importe"

    if nuevas_cols:
        df.rename(columns=nuevas_cols, inplace=True)
        print(f"‚úÖ [BANCO] Columnas renombradas: {nuevas_cols}")
    else:
        # Esta rama deber√≠a ser rara porque arriba devolvemos df si mapeo <2.
        print("‚ö†Ô∏è [BANCO] No se pudo mapear autom√°ticamente (fallback)")

    for c in ["Fecha", "Concepto", "Importe"]:
        if c not in df.columns:
            print(f"‚ö†Ô∏è [BANCO] Columna faltante: {c}, rellenando con None")
            df[c] = None

    if "Fecha" in df.columns:
        df["Fecha"] = pd.to_datetime(df["Fecha"], errors="coerce", dayfirst=True)

    if "Importe" in df.columns:
        df["Importe"] = (
            df["Importe"]
            .astype(str)
            .str.replace("‚Ç¨", "")
            .str.replace(".", "")
            .str.replace(",", ".")
        )
        df["Importe"] = pd.to_numeric(df["Importe"], errors="coerce").fillna(0)

    print(f"‚úÖ [BANCO] Normalizaci√≥n completada ({len(df)} filas)")
    return df[["Fecha", "Concepto", "Importe"]]


# =========================================================
# INGESTA EXCEL (FASE 1C: MULTI-FORMATO)
# =========================================================


def leer_excel(file_stream, filename: str) -> ParsingResult:
    """
    Lee un archivo Excel (.xlsx, .xls) y extrae su contenido como texto estructurado.

    FASE 1C: Parser dedicado para Excel.

    Diferencias con leer_csv_excel:
    - Extrae texto, no datos estructurados
    - Preserva formato de hojas y celdas
    - Retorna ParsingResult (compatible con chunking)
    - Incluye offsets de hojas para trazabilidad

    Args:
        file_stream: Stream del archivo Excel
        filename: Nombre del archivo (para logging)

    Returns:
        ParsingResult con texto extra√≠do y metadatos

    Casos de uso:
        - Balances de situaci√≥n
        - Cuentas de PyG
        - Listados de acreedores
        - Extractos bancarios
    """
    print(f"üìä [EXCEL] Inicio lectura de Excel: {filename}")

    try:
        # Convertir stream a BytesIO si es necesario
        if isinstance(file_stream, bytes):
            file_stream = io.BytesIO(file_stream)
        elif isinstance(file_stream, (str, Path)):
            # Si es ruta, abrir archivo
            with open(file_stream, "rb") as f:
                file_stream = io.BytesIO(f.read())

        # Usar el nuevo parser de Excel
        excel_result: ExcelParseResult = parse_excel_stream(file_stream, filename)

        print("‚úÖ [EXCEL] Excel le√≠do exitosamente")
        print(f"‚úÖ [EXCEL] - Hojas: {excel_result.num_paginas}")
        print(f"‚úÖ [EXCEL] - Caracteres: {len(excel_result.texto)}")

        # FASE 2A: Detectar si es un estado financiero
        structured_data = None
        from app.services.accounting_parser import (
            is_financial_statement,
            parse_financial_statement_from_text,
        )

        if is_financial_statement(excel_result.texto):
            print("üìä [EXCEL] Detectado estado financiero, extrayendo datos...")
            financial_stmt = parse_financial_statement_from_text(excel_result.texto)
            if financial_stmt:
                structured_data = {"financial_statement": financial_stmt.model_dump()}

        # Convertir ExcelParseResult a ParsingResult (compatible con sistema existente)
        return ParsingResult(
            texto=excel_result.texto,
            num_paginas=excel_result.num_paginas,
            tipo_documento="excel",
            page_offsets=excel_result.page_offsets,
            structured_data=structured_data,
        )

    except Exception as e:
        print("‚ùå [EXCEL] Error leyendo Excel")
        print(f"‚ùå [EXCEL] Detalle: {e}")

        # Retornar ParsingResult vac√≠o en caso de error
        return ParsingResult(
            texto="",
            num_paginas=0,
            tipo_documento="excel",
            page_offsets={},
        )


# =========================================================
# INGESTA WORD (FASE 1D: MULTI-FORMATO)
# =========================================================


def leer_word(file_stream, filename: str) -> ParsingResult:
    """
    Lee un archivo Word (.docx) y extrae su contenido como texto estructurado.

    FASE 1D: Parser dedicado para Word (reemplaza leer_docx legacy).

    Diferencias con leer_docx existente:
    - Extrae p√°rrafos Y tablas de forma estructurada
    - Detecta estilos de encabezados
    - Estima p√°ginas por n√∫mero de p√°rrafos
    - Incluye offsets estimados para trazabilidad
    - Retorna ParsingResult (compatible con chunking)

    Args:
        file_stream: Stream del archivo Word
        filename: Nombre del archivo (para logging)

    Returns:
        ParsingResult con texto extra√≠do y metadatos

    Casos de uso:
        - Informes previos de auditor√≠a
        - Contratos mercantiles
        - Escritos de acreedores
        - Comunicaciones formales
        - Documentos societarios
    """
    print(f"üìù [WORD] Inicio lectura de Word: {filename}")

    try:
        # Convertir stream a BytesIO si es necesario
        if isinstance(file_stream, bytes):
            file_stream = io.BytesIO(file_stream)
        elif isinstance(file_stream, (str, Path)):
            # Si es ruta, abrir archivo
            with open(file_stream, "rb") as f:
                file_stream = io.BytesIO(f.read())

        # Usar el nuevo parser de Word
        word_result: WordParseResult = parse_word_stream(file_stream, filename)

        print("‚úÖ [WORD] Word le√≠do exitosamente")
        print(f"‚úÖ [WORD] - P√°ginas estimadas: {word_result.num_paginas}")
        print(f"‚úÖ [WORD] - P√°rrafos: {len(word_result.paragraphs_content)}")
        print(f"‚úÖ [WORD] - Tablas: {len(word_result.tables_content)}")
        print(f"‚úÖ [WORD] - Caracteres: {len(word_result.texto)}")

        # Convertir WordParseResult a ParsingResult (compatible con sistema existente)
        return ParsingResult(
            texto=word_result.texto,
            num_paginas=word_result.num_paginas,
            tipo_documento="word",
            page_offsets=word_result.page_offsets,
        )

    except Exception as e:
        print("‚ùå [WORD] Error leyendo Word")
        print(f"‚ùå [WORD] Detalle: {e}")

        # Retornar ParsingResult vac√≠o en caso de error
        return ParsingResult(
            texto="",
            num_paginas=0,
            tipo_documento="word",
            page_offsets={},
        )


def leer_csv_excel(file_stream, filename: str) -> Optional[pd.DataFrame]:
    import pandas as pd

    print(f"üìä [CSV/EXCEL] Procesando archivo: {filename}")

    try:
        name = filename.lower()

        if name.endswith(".csv"):
            print("üìä [CSV] Detectado CSV")
            df = pd.read_csv(file_stream)
        elif name.endswith((".xls", ".xlsx")):
            print("üìä [EXCEL] Detectado Excel")
            df = pd.read_excel(file_stream)
        else:
            print("‚ö†Ô∏è [CSV/EXCEL] Formato no compatible")
            return None

        print(f"üìä [CSV/EXCEL] DataFrame cargado ({df.shape[0]} filas)")
        return normalizar_datos_banco(df)

    except Exception as e:
        print("‚ùå [CSV/EXCEL] Error leyendo archivo")
        print(f"‚ùå [CSV/EXCEL] Detalle: {e}")
        return None


# =========================================================
# FUNCI√ìN P√öBLICA √öNICA (DISPATCHER)
# =========================================================


def ingerir_archivo(
    file_stream,
    filename: str,
    file_path: Optional[Path] = None,
    case_id: Optional[str] = None,
    validation_mode: Optional[str] = None,
) -> Union[ParsingResult, pd.DataFrame, None]:
    """
    Punto √∫nico de entrada para ingesta.
    Detecta formato y delega la lectura.

    ENDURECIMIENTO #2 (FASE 3): Validaci√≥n FAIL-FAST integrada.
    Si file_path, case_id y validation_mode son provistos, ejecuta validaci√≥n ANTES de parsing.

    Args:
        file_stream: Stream del archivo
        filename: Nombre del archivo
        file_path: Path del archivo (requerido para validaci√≥n)
        case_id: ID del caso (requerido para validaci√≥n)
        validation_mode: "strict" o "permissive" (requerido para validaci√≥n)

    Retorna:
    - ParsingResult para archivos de texto (PDF, TXT, DOCX, DOC)
    - DataFrame para CSV/Excel
    - None si el formato no es soportado o si validaci√≥n falla
    """
    from app.services.ingestion_failfast import (
        ValidationMode,
        should_skip_document,
        validate_document_failfast,
    )

    print("--------------------------------------------------")
    print(f"üì• [INGESTA] Archivo recibido: {filename}")

    # FAIL-FAST: Validaci√≥n PRE-ingesta (si par√°metros provistos)
    if file_path and case_id and validation_mode:
        print(f"üì• [INGESTA] Ejecutando validaci√≥n FAIL-FAST (modo={validation_mode})")

        mode = ValidationMode.STRICT if validation_mode == "strict" else ValidationMode.PERMISSIVE

        validation_result = validate_document_failfast(
            file_path=file_path,
            extracted_text=None,
            case_id=case_id,
            mode=mode,
        )

        if should_skip_document(validation_result):
            print(f"‚ùå [INGESTA] Documento rechazado en validaci√≥n: {filename}")
            print(f"‚ùå [INGESTA] Raz√≥n: {validation_result.reject_code}")
            return None

    name = filename.lower()

    if name.endswith(".pdf"):
        print("üì• [INGESTA] Tipo detectado: PDF")
        return leer_pdf(file_stream)

    if name.endswith(".txt"):
        print("üì• [INGESTA] Tipo detectado: TXT")
        return leer_txt(file_stream)

    # FASE 1D: Soporte mejorado para Word
    if name.endswith(".docx"):
        print("üì• [INGESTA] Tipo detectado: DOCX")
        return leer_word(file_stream, filename)

    if name.endswith(".doc"):
        print("üì• [INGESTA] Tipo detectado: DOC (legacy, best effort)")
        # Intentar con el nuevo parser (python-docx soporta algunos .doc)
        try:
            return leer_word(file_stream, filename)
        except Exception as e:
            print(f"‚ö†Ô∏è [INGESTA] Fall√≥ parser moderno, intentando legacy: {e}")
            # Fallback al parser antiguo
            return leer_docx(file_stream, is_doc_legacy=True)

    # FASE 1C: Soporte espec√≠fico para Excel
    if name.endswith((".xlsx", ".xls")):
        print("üì• [INGESTA] Tipo detectado: EXCEL (.xlsx/.xls)")
        return leer_excel(file_stream, filename)

    if name.endswith(".csv"):
        print("üì• [INGESTA] Tipo detectado: CSV")
        return leer_csv_excel(file_stream, filename)

    # FASE 2A: Soporte para emails
    if name.endswith(".eml"):
        print("üì• [INGESTA] Tipo detectado: EMAIL (.eml)")
        result = parse_eml_stream(file_stream)
        # Convertir EmailParseResult a ParsingResult
        return ParsingResult(
            texto=result.texto,
            num_paginas=result.num_paginas,
            tipo_documento=result.tipo_documento,
            page_offsets=result.page_offsets,
        )

    if name.endswith(".msg"):
        print("üì• [INGESTA] Tipo detectado: EMAIL (.msg - Outlook)")
        result = parse_msg_stream(file_stream, filename)
        # Convertir EmailParseResult a ParsingResult
        return ParsingResult(
            texto=result.texto,
            num_paginas=result.num_paginas,
            tipo_documento=result.tipo_documento,
            page_offsets=result.page_offsets,
        )

    # FASE 2A: Soporte para im√°genes con OCR
    if name.endswith((".jpg", ".jpeg", ".png", ".tiff", ".tif")):
        print(f"üì• [INGESTA] Tipo detectado: IMAGEN ({name.split('.')[-1].upper()})")

        if not is_tesseract_available():
            print("‚ùå [INGESTA] Tesseract no disponible, no se puede procesar imagen")
            return ParsingResult(
                texto="ERROR: Tesseract no instalado",
                num_paginas=0,
                tipo_documento="image",
                page_offsets=None,
            )

        # Leer bytes de la imagen
        if hasattr(file_stream, "read"):
            image_bytes = file_stream.read()
        else:
            with open(file_stream, "rb") as f:
                image_bytes = f.read()

        # Aplicar OCR
        from app.services.ocr_parser import ocr_image_from_bytes

        ocr_result = ocr_image_from_bytes(image_bytes, lang="spa")

        return ParsingResult(
            texto=ocr_result.texto,
            num_paginas=ocr_result.num_paginas,
            tipo_documento="image",
            page_offsets=ocr_result.page_offsets,
        )

    print(f"‚ö†Ô∏è [INGESTA] Formato no soportado: {filename}")
    return None
