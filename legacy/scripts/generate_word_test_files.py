"""
Generador de archivos Word de prueba para RETAIL DEMO SL.

Genera:
- Informe de auditoría previo
- Contrato de préstamo mercantil

Con datos coherentes para probar el parser de Word.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT_DIR = Path("data/casos_prueba/RETAIL_DEMO_SL")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_audit_report():
    """Genera informe de auditoría en formato Word."""
    doc = Document()

    # Configurar márgenes y estilos
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Título
    title = doc.add_heading("INFORME DE AUDITORÍA INTERNA", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo
    subtitle = doc.add_paragraph("RETAIL DEMO SL")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True

    subtitle2 = doc.add_paragraph("Ejercicio 2023")
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_run = subtitle2.runs[0]
    subtitle2_run.font.italic = True

    doc.add_paragraph()  # Espacio

    # 1. INTRODUCCIÓN
    doc.add_heading("1. INTRODUCCIÓN", level=2)

    doc.add_paragraph(
        "El presente informe tiene por objeto analizar la situación económico-financiera de "
        "RETAIL DEMO SL (en adelante, 'la Sociedad') a 31 de diciembre de 2023, con el fin de "
        "determinar si concurren las circunstancias previstas en el artículo 2.2 de la Ley Concursal "
        "que obligarían a la solicitud de declaración de concurso de acreedores."
    )

    # 2. SITUACIÓN PATRIMONIAL
    doc.add_heading("2. SITUACIÓN PATRIMONIAL", level=2)

    doc.add_paragraph(
        "El análisis del balance de situación a 31 de diciembre de 2023 revela una situación de "
        "insolvencia actual, según se detalla a continuación:"
    )

    # Tabla de balance resumido
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"

    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "CONCEPTO"
    hdr_cells[1].text = "IMPORTE (€)"

    # Datos
    table.rows[1].cells[0].text = "Activo Total"
    table.rows[1].cells[1].text = "250.000"

    table.rows[2].cells[0].text = "Pasivo Total"
    table.rows[2].cells[1].text = "480.000"

    table.rows[3].cells[0].text = "Patrimonio Neto"
    table.rows[3].cells[1].text = "-230.000"

    table.rows[4].cells[0].text = "Ratio de Solvencia"
    table.rows[4].cells[1].text = "0.52"

    doc.add_paragraph()

    conclusion_pat = doc.add_paragraph(
        "CONCLUSIÓN: El patrimonio neto es negativo por importe de 230.000€, lo que implica que "
        "el pasivo exigible supera el valor del activo. Esta situación constituye insolvencia actual "
        "en los términos del artículo 2.2 de la Ley Concursal."
    )
    conclusion_pat.runs[0].font.bold = True

    # 3. ANÁLISIS DE RESULTADOS
    doc.add_heading("3. ANÁLISIS DE RESULTADOS", level=2)

    doc.add_paragraph(
        "La cuenta de pérdidas y ganancias del ejercicio 2023 muestra una evolución negativa "
        "de la actividad:"
    )

    # Tabla de PyG resumida
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = "Light Grid Accent 1"

    # Encabezados
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = "CONCEPTO"
    hdr_cells2[1].text = "IMPORTE (€)"

    # Datos
    table2.rows[1].cells[0].text = "Ingresos de explotación"
    table2.rows[1].cells[1].text = "120.000"

    table2.rows[2].cells[0].text = "Gastos de explotación"
    table2.rows[2].cells[1].text = "-183.000"

    table2.rows[3].cells[0].text = "Resultado de explotación"
    table2.rows[3].cells[1].text = "-63.000"

    table2.rows[4].cells[0].text = "Resultado financiero"
    table2.rows[4].cells[1].text = "-12.000"

    table2.rows[5].cells[0].text = "RESULTADO DEL EJERCICIO"
    table2.rows[5].cells[1].text = "-60.000"

    doc.add_paragraph()

    doc.add_paragraph(
        "Las pérdidas del ejercicio ascienden a 60.000€, agravando la situación patrimonial "
        "ya comprometida de ejercicios anteriores."
    )

    # 4. ANÁLISIS DE LIQUIDEZ
    doc.add_heading("4. ANÁLISIS DE LIQUIDEZ", level=2)

    doc.add_paragraph("El análisis de liquidez revela una situación crítica:")

    liquidez = doc.add_paragraph()
    liquidez.add_run("• ").bold = True
    liquidez.add_run("Activo corriente: 70.000€\n")
    liquidez.add_run("• ").bold = True
    liquidez.add_run("Pasivo corriente: 300.000€\n")
    liquidez.add_run("• ").bold = True
    liquidez.add_run("Ratio de liquidez: 0.23\n")
    liquidez.add_run("• ").bold = True
    liquidez.add_run("Efectivo disponible: 3.500€")

    doc.add_paragraph(
        "La sociedad no dispone de recursos suficientes para atender sus obligaciones de pago "
        "a corto plazo. El efectivo disponible es insuficiente para cubrir la deuda con Hacienda "
        "(68.000€) y Seguridad Social (42.000€)."
    )

    # 5. CONCLUSIONES Y RECOMENDACIONES
    doc.add_heading("5. CONCLUSIONES Y RECOMENDACIONES", level=2)

    conclusiones = doc.add_paragraph()
    conclusiones.add_run("A) INSOLVENCIA ACTUAL\n").bold = True
    conclusiones.add_run(
        "La sociedad se encuentra en situación de insolvencia actual conforme al artículo 2.2 "
        "de la Ley Concursal, al resultar que el valor de sus activos es inferior al de sus pasivos.\n\n"
    )

    conclusiones.add_run("B) OBLIGACIÓN DE SOLICITAR CONCURSO\n").bold = True
    conclusiones.add_run(
        "De conformidad con el artículo 5 de la Ley Concursal, el deudor deberá solicitar la "
        "declaración de concurso dentro de los dos meses siguientes a la fecha en que hubiera "
        "conocido o debido conocer su estado de insolvencia.\n\n"
    )

    conclusiones.add_run("C) RIESGO DE CALIFICACIÓN CULPABLE\n").bold = True
    conclusiones.add_run(
        "El retraso en la solicitud de concurso, así como la realización de pagos a personas "
        "especialmente relacionadas con el deudor durante los dos años anteriores, podrían dar "
        "lugar a una calificación culpable del concurso (arts. 164 y 257 LC)."
    )

    doc.add_paragraph()

    # Firma
    firma = doc.add_paragraph("Madrid, 15 de diciembre de 2023")
    firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    firma_auditor = doc.add_paragraph("Fdo.: AUDITORÍA Y CONSULTORÍA SL")
    firma_auditor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    firma_auditor_run = firma_auditor.runs[0]
    firma_auditor_run.font.italic = True

    # Guardar
    filename = OUTPUT_DIR / "12_Informe_Auditoria_2023.docx"
    doc.save(filename)
    print(f"✅ Creado: {filename}")


def create_loan_contract():
    """Genera contrato de préstamo en formato Word."""
    doc = Document()

    # Título
    title = doc.add_heading("CONTRATO DE PRÉSTAMO MERCANTIL", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Partes
    doc.add_heading("PARTES CONTRATANTES", level=2)

    doc.add_paragraph("En Madrid, a 15 de enero de 2022, COMPARECEN:")

    doc.add_paragraph()

    parte_a = doc.add_paragraph()
    parte_a.add_run("DE UNA PARTE, ").bold = True
    parte_a.add_run(
        "D. Juan García López, mayor de edad, con DNI 12345678A, en nombre y representación de "
        "BANCO EJEMPLO SA, con CIF A-12345678, con domicilio social en Madrid, calle Gran Vía 123, "
        "en su calidad de Director de Banca Empresas, según consta en escritura pública de poderes "
        "otorgada ante el Notario de Madrid D. Antonio Martínez el 10 de enero de 2020."
    )

    doc.add_paragraph()

    parte_b = doc.add_paragraph()
    parte_b.add_run("DE OTRA PARTE, ").bold = True
    parte_b.add_run(
        "D. Pedro Sánchez Ruiz, mayor de edad, con DNI 87654321B, en nombre y representación de "
        "RETAIL DEMO SL, con CIF B-87654321, con domicilio social en Madrid, calle Mayor 123, "
        "en su calidad de Administrador Único."
    )

    doc.add_paragraph()

    # Antecedentes
    doc.add_heading("ANTECEDENTES", level=2)

    doc.add_paragraph(
        "I. Que RETAIL DEMO SL necesita financiación para el desarrollo de su actividad empresarial."
    )

    doc.add_paragraph(
        "II. Que BANCO EJEMPLO SA está dispuesto a conceder un préstamo a RETAIL DEMO SL en las "
        "condiciones que se establecen en el presente contrato."
    )

    doc.add_paragraph()

    # Estipulaciones
    doc.add_heading("ESTIPULACIONES", level=2)

    doc.add_heading("PRIMERA. Objeto del contrato", level=3)
    doc.add_paragraph(
        "BANCO EJEMPLO SA (en adelante, 'EL PRESTAMISTA') concede a RETAIL DEMO SL (en adelante, "
        "'EL PRESTATARIO') un préstamo mercantil por importe de CIENTO OCHENTA MIL EUROS (180.000€)."
    )

    doc.add_heading("SEGUNDA. Plazo y amortización", level=3)
    doc.add_paragraph(
        "El préstamo se concede por un plazo de 5 años, a contar desde la fecha de disposición. "
        "La amortización se realizará mediante cuotas mensuales constantes de 3.200€, comprensivas "
        "de capital e intereses."
    )

    doc.add_heading("TERCERA. Tipo de interés", level=3)
    doc.add_paragraph(
        "El tipo de interés aplicable será del 4,5% nominal anual, calculado sobre el capital "
        "pendiente de amortización."
    )

    doc.add_heading("CUARTA. Destino del préstamo", level=3)
    doc.add_paragraph(
        "EL PRESTATARIO destinará el importe del préstamo exclusivamente a financiar:"
    )
    dest = doc.add_paragraph()
    dest.add_run("• ").bold = True
    dest.add_run("Adquisición de maquinaria: 120.000€\n")
    dest.add_run("• ").bold = True
    dest.add_run("Reforma de instalaciones: 60.000€")

    doc.add_heading("QUINTA. Garantías", level=3)
    doc.add_paragraph(
        "Como garantía del cumplimiento de las obligaciones derivadas del presente contrato, "
        "EL PRESTATARIO constituye hipoteca sobre el inmueble sito en Madrid, calle Mayor 123, "
        "inscrito en el Registro de la Propiedad nº 5 de Madrid, al tomo 1234, libro 567, folio 89."
    )

    doc.add_heading("SEXTA. Vencimiento anticipado", level=3)
    doc.add_paragraph(
        "EL PRESTAMISTA podrá declarar vencido anticipadamente el préstamo y exigir el reembolso "
        "total del capital pendiente más los intereses devengados en los siguientes supuestos:"
    )
    venc = doc.add_paragraph()
    venc.add_run("a) ").bold = True
    venc.add_run("Impago de dos cuotas consecutivas.\n")
    venc.add_run("b) ").bold = True
    venc.add_run("Declaración de concurso de acreedores.\n")
    venc.add_run("c) ").bold = True
    venc.add_run("Incumplimiento del destino del préstamo.")

    doc.add_heading("SÉPTIMA. Jurisdicción", level=3)
    doc.add_paragraph(
        "Para cualquier cuestión litigiosa que pudiera derivarse del presente contrato, las partes "
        "se someten expresamente a los Juzgados y Tribunales de Madrid, con renuncia a cualquier "
        "otro fuero que pudiera corresponderles."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Firmas
    doc.add_paragraph(
        "Y en prueba de conformidad, firman el presente contrato en el lugar y "
        "fecha indicados en el encabezamiento."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabla de firmas
    firma_table = doc.add_table(rows=1, cols=2)
    firma_cells = firma_table.rows[0].cells

    firma_cells[0].text = "EL PRESTAMISTA\n\n\n\nBanco Ejemplo SA"
    firma_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    firma_cells[1].text = "EL PRESTATARIO\n\n\n\nRETAIL DEMO SL"
    firma_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar
    filename = OUTPUT_DIR / "13_Contrato_Prestamo_BancoEjemplo.docx"
    doc.save(filename)
    print(f"✅ Creado: {filename}")


def main():
    """Genera todos los archivos Word de prueba."""
    print("\n🚀 Generando archivos Word para RETAIL DEMO SL\n")
    print("=" * 60)

    create_audit_report()
    create_loan_contract()

    print("=" * 60)
    print("\n✅ ARCHIVOS WORD GENERADOS\n")
    print(f"📁 Ubicación: {OUTPUT_DIR}")
    print("📄 Archivos: 2 documentos Word")
    print("\n🎯 Listos para probar el parser de Word\n")


if __name__ == "__main__":
    main()
